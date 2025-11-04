"""
KUMPAS Kalibrasyon Sertifikası Word Generator
AS KALİBRASYON formatında profesyonel Word belgesi oluşturur
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional


def create_kumpas_certificate(
    kalibrasyon_data: Dict,
    cihaz_data: Dict,
    organizasyon_data: Dict,
    output_dir: Path
) -> str:
    """
    KUMPAS kalibrasyon sertifikası Word belgesi oluşturur
    
    Args:
        kalibrasyon_data: Kalibrasyon verileri (ölçümler, sapma, vb)
        cihaz_data: Cihaz bilgileri (marka, model, seri no)
        organizasyon_data: Müşteri/organizasyon bilgileri
        output_dir: Word dosyasının kaydedileceği dizin
    
    Returns:
        str: Oluşturulan dosya yolu
    """
    
    doc = Document()
    
    # Sayfa ayarları (A4)
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    
    # Sertifika numarası oluştur
    sertifika_no = f"AS-{datetime.now().strftime('%Y%m%d-%H%M%S')}"
    
    # ================== BAŞLIK BÖLGESİ ==================
    # Logo ve firma bilgisi için tablo
    header_table = doc.add_table(rows=1, cols=2)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Sol: Logo alanı
    cell_logo = header_table.rows[0].cells[0]
    p_logo = cell_logo.paragraphs[0]
    p_logo.text = "AS KALİBRASYON"
    p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_logo = p_logo.runs[0]
    run_logo.font.size = Pt(18)
    run_logo.font.bold = True
    run_logo.font.color.rgb = RGBColor(0, 51, 153)
    
    # Sağ: İletişim bilgileri
    cell_contact = header_table.rows[0].cells[1]
    p_contact = cell_contact.paragraphs[0]
    p_contact.text = "Telefon: +90 XXX XXX XX XX\nE-posta: info@askalibrasyon.com\nWeb: www.askalibrasyon.com"
    p_contact.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_contact = p_contact.runs[0]
    run_contact.font.size = Pt(9)
    
    # Boşluk
    doc.add_paragraph()
    
    # ================== SERTİFİKA BAŞLIĞI ==================
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = title.add_run("KALİBRASYON SERTİFİKASI")
    run_title.font.size = Pt(16)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0, 51, 153)
    
    # Sertifika numarası
    cert_no = doc.add_paragraph()
    cert_no.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_cert = cert_no.add_run(f"Sertifika No: {sertifika_no}")
    run_cert.font.size = Pt(11)
    run_cert.font.bold = True
    
    # Boşluk
    doc.add_paragraph()
    
    # ================== MÜŞTERİ BİLGİLERİ ==================
    info_table = doc.add_table(rows=8, cols=2)
    info_table.style = 'Table Grid'
    
    # Başlıklar ve değerler
    info_data = [
        ("Müşteri Adı", organizasyon_data.get("musteri_adi", "N/A")),
        ("Müşteri Adresi", organizasyon_data.get("adres", "N/A")),
        ("İstek No", kalibrasyon_data.get("talep_no", "N/A")),
        ("Cihaz Tipi", "Kumpas (Dijital/Analog)"),
        ("Marka", cihaz_data.get("marka", "N/A")),
        ("Model", cihaz_data.get("model", "N/A")),
        ("Seri No", cihaz_data.get("seri_no", "N/A")),
        ("Ölçme Aralığı", cihaz_data.get("olcme_araligi", "0-150 mm")),
    ]
    
    for i, (label, value) in enumerate(info_data):
        row = info_table.rows[i]
        
        # Başlık hücresi
        cell_label = row.cells[0]
        cell_label.paragraphs[0].text = label
        cell_label.paragraphs[0].runs[0].font.bold = True
        cell_label.paragraphs[0].runs[0].font.size = Pt(10)
        cell_label.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # Değer hücresi
        cell_value = row.cells[1]
        cell_value.paragraphs[0].text = str(value)
        cell_value.paragraphs[0].runs[0].font.size = Pt(10)
        cell_value.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    # Boşluk
    doc.add_paragraph()
    
    # ================== ÇEVRE KOŞULLARI ==================
    env_heading = doc.add_paragraph()
    run_env = env_heading.add_run("Kalibrasyon Çevre Koşulları")
    run_env.font.size = Pt(12)
    run_env.font.bold = True
    run_env.font.color.rgb = RGBColor(0, 51, 153)
    
    env_table = doc.add_table(rows=3, cols=2)
    env_table.style = 'Table Grid'
    
    env_data = [
        ("Kalibrasyon Tarihi", datetime.now().strftime("%d.%m.%Y")),
        ("Sıcaklık", f"{kalibrasyon_data.get('sicaklik', 'N/A')} °C"),
        ("Bağıl Nem", f"{kalibrasyon_data.get('nem', 'N/A')} %"),
    ]
    
    for i, (label, value) in enumerate(env_data):
        row = env_table.rows[i]
        cell_label = row.cells[0]
        cell_label.paragraphs[0].text = label
        cell_label.paragraphs[0].runs[0].font.bold = True
        cell_label.paragraphs[0].runs[0].font.size = Pt(10)
        
        cell_value = row.cells[1]
        cell_value.paragraphs[0].text = str(value)
        cell_value.paragraphs[0].runs[0].font.size = Pt(10)
    
    # Boşluk
    doc.add_paragraph()
    
    # ================== ÖLÇÜM SONUÇLARI ==================
    results_heading = doc.add_paragraph()
    run_results = results_heading.add_run("Ölçüm Sonuçları")
    run_results.font.size = Pt(12)
    run_results.font.bold = True
    run_results.font.color.rgb = RGBColor(0, 51, 153)
    
    # DIŞ ÇENE ÖLÇÜMLERİ
    if kalibrasyon_data.get("olcum_verileri", {}).get("dis_olcum"):
        doc.add_paragraph("A) Dış Çene Ölçümleri", style='Heading 3')
        
        dis_olcum = kalibrasyon_data["olcum_verileri"]["dis_olcum"]
        dis_table = doc.add_table(rows=len(dis_olcum) + 1, cols=4)
        dis_table.style = 'Table Grid'
        
        # Tablo başlıkları
        headers = ["Nominal (mm)", "Ölçülen (mm)", "Sapma (mm)", "U (mm)"]
        for j, header in enumerate(headers):
            cell = dis_table.rows[0].cells[j]
            cell.paragraphs[0].text = header
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Başlık arka plan rengi (mavi)
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), "4472C4")
            cell._element.get_or_add_tcPr().append(shading_elm)
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        
        # Veri satırları
        for i, olcum in enumerate(dis_olcum, 1):
            row = dis_table.rows[i]
            
            # Nominal
            row.cells[0].text = f"{olcum['referans_deger']:.2f}"
            row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Ölçülen
            row.cells[1].text = f"{olcum['okunan_deger']:.2f}"
            row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Sapma
            sapma = olcum.get('sapma', 0)
            row.cells[2].text = f"{sapma:.3f}"
            row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Sapma renklendir
            if abs(sapma) > 0.02:
                # Kırmızı (tolerans dışı)
                row.cells[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)
                row.cells[2].paragraphs[0].runs[0].font.bold = True
            elif sapma != 0:
                # Turuncu (tolerans içinde)
                row.cells[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 165, 0)
            else:
                # Yeşil (sıfır sapma)
                row.cells[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 128, 0)
            
            # Belirsizlik
            row.cells[3].text = f"{olcum.get('belirsizlik', 0.03):.2f}"
            row.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # İÇ ÇENE ÖLÇÜMLERİ
    if kalibrasyon_data.get("olcum_verileri", {}).get("ic_olcum"):
        doc.add_paragraph()
        doc.add_paragraph("B) İç Çene Ölçümleri", style='Heading 3')
        
        ic_olcum = kalibrasyon_data["olcum_verileri"]["ic_olcum"]
        ic_table = doc.add_table(rows=len(ic_olcum) + 1, cols=4)
        ic_table.style = 'Table Grid'
        
        # Tablo başlıkları
        for j, header in enumerate(headers):
            cell = ic_table.rows[0].cells[j]
            cell.paragraphs[0].text = header
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), "4472C4")
            cell._element.get_or_add_tcPr().append(shading_elm)
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        
        # Veri satırları
        for i, olcum in enumerate(ic_olcum, 1):
            row = ic_table.rows[i]
            row.cells[0].text = f"{olcum['referans_deger']:.2f}"
            row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row.cells[1].text = f"{olcum['okunan_deger']:.2f}"
            row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            sapma = olcum.get('sapma', 0)
            row.cells[2].text = f"{sapma:.3f}"
            row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            if abs(sapma) > 0.02:
                row.cells[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)
                row.cells[2].paragraphs[0].runs[0].font.bold = True
            elif sapma != 0:
                row.cells[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 165, 0)
            else:
                row.cells[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 128, 0)
            
            row.cells[3].text = f"{olcum.get('belirsizlik', 0.03):.2f}"
            row.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # DERİNLİK ÖLÇÜMLERİ
    if kalibrasyon_data.get("olcum_verileri", {}).get("derinlik_olcum"):
        doc.add_paragraph()
        doc.add_paragraph("C) Derinlik Ölçümleri", style='Heading 3')
        
        derinlik_olcum = kalibrasyon_data["olcum_verileri"]["derinlik_olcum"]
        depth_table = doc.add_table(rows=len(derinlik_olcum) + 1, cols=4)
        depth_table.style = 'Table Grid'
        
        for j, header in enumerate(headers):
            cell = depth_table.rows[0].cells[j]
            cell.paragraphs[0].text = header
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), "4472C4")
            cell._element.get_or_add_tcPr().append(shading_elm)
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        
        for i, olcum in enumerate(derinlik_olcum, 1):
            row = depth_table.rows[i]
            row.cells[0].text = f"{olcum['referans_deger']:.2f}"
            row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row.cells[1].text = f"{olcum['okunan_deger']:.2f}"
            row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            sapma = olcum.get('sapma', 0)
            row.cells[2].text = f"{sapma:.3f}"
            row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            if abs(sapma) > 0.02:
                row.cells[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)
                row.cells[2].paragraphs[0].runs[0].font.bold = True
            elif sapma != 0:
                row.cells[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 165, 0)
            else:
                row.cells[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 128, 0)
            
            row.cells[3].text = f"{olcum.get('belirsizlik', 0.03):.2f}"
            row.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Boşluk
    doc.add_paragraph()
    
    # ================== FONKSİYONELLİK KONTROLLARI ==================
    func_heading = doc.add_paragraph()
    run_func = func_heading.add_run("Fonksiyonellik Kontrolleri")
    run_func.font.size = Pt(12)
    run_func.font.bold = True
    run_func.font.color.rgb = RGBColor(0, 51, 153)
    
    fonk_data = kalibrasyon_data.get("olcum_verileri", {}).get("fonksiyonellik", {})
    
    func_table = doc.add_table(rows=5, cols=2)
    func_table.style = 'Table Grid'
    
    func_items = [
        ("Ölçme Çeneleri", fonk_data.get("olcme_ceneleri", "N/A")),
        ("Tespit Vidası", fonk_data.get("tespit_vidasi", "N/A")),
        ("Gösterge", fonk_data.get("gosterge", "N/A")),
        ("Tambur/Vernier", fonk_data.get("tambur_vernier", "N/A")),
    ]
    
    # Başlık satırı
    func_table.rows[0].cells[0].paragraphs[0].text = "Kontrol Noktası"
    func_table.rows[0].cells[0].paragraphs[0].runs[0].font.bold = True
    func_table.rows[0].cells[1].paragraphs[0].text = "Durum"
    func_table.rows[0].cells[1].paragraphs[0].runs[0].font.bold = True
    
    for i, (kontrol, durum) in enumerate(func_items, 1):
        func_table.rows[i].cells[0].text = kontrol
        func_table.rows[i].cells[1].text = durum
        
        # Durum renklendirme
        if durum == "Uygun":
            func_table.rows[i].cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 128, 0)
            func_table.rows[i].cells[1].paragraphs[0].runs[0].font.bold = True
        elif durum == "Uygun Değil":
            func_table.rows[i].cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)
            func_table.rows[i].cells[1].paragraphs[0].runs[0].font.bold = True
    
    # Boşluk
    doc.add_paragraph()
    
    # ================== REFERANS STANDARTLAR ==================
    ref_heading = doc.add_paragraph()
    run_ref = ref_heading.add_run("Referans Standartlar ve Cihazlar")
    run_ref.font.size = Pt(12)
    run_ref.font.bold = True
    run_ref.font.color.rgb = RGBColor(0, 51, 153)
    
    ref_standards = [
        "MEK.SIT.002 - Kumpas Kalibrasyon Prosedürü",
        "ISO/IEC 17025:2017",
        "Referans Blok Set: Sertifika No: AS-REF-2024-001",
    ]
    
    for standard in ref_standards:
        p_std = doc.add_paragraph(standard, style='List Bullet')
        p_std.runs[0].font.size = Pt(10)
    
    # Boşluk
    doc.add_paragraph()
    
    # ================== SONUÇ ==================
    uygunluk = kalibrasyon_data.get("uygunluk", False)
    
    conclusion = doc.add_paragraph()
    run_conclusion = conclusion.add_run("Sonuç: ")
    run_conclusion.font.size = Pt(12)
    run_conclusion.font.bold = True
    
    result_text = "UYGUN" if uygunluk else "UYGUN DEĞİL"
    run_result = conclusion.add_run(result_text)
    run_result.font.size = Pt(12)
    run_result.font.bold = True
    
    if uygunluk:
        run_result.font.color.rgb = RGBColor(0, 128, 0)
    else:
        run_result.font.color.rgb = RGBColor(255, 0, 0)
    
    # Boşluk
    doc.add_paragraph()
    doc.add_paragraph()
    
    # ================== İMZA ALANLARI ==================
    sig_table = doc.add_table(rows=3, cols=3)
    
    # Başlıklar
    sig_table.rows[0].cells[0].text = "Kalibrasyonu Yapan"
    sig_table.rows[0].cells[1].text = "Kontrol Eden"
    sig_table.rows[0].cells[2].text = "Onaylayan"
    
    for cell in sig_table.rows[0].cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # İmza çizgileri
    for i in range(3):
        sig_table.rows[1].cells[i].text = "\n\n_________________"
        sig_table.rows[1].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # İsimler (boş - elle doldurulacak)
    for i in range(3):
        sig_table.rows[2].cells[i].text = "\n"
        sig_table.rows[2].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Boşluk
    doc.add_paragraph()
    
    # ================== FOOTER ==================
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_footer = footer_p.add_run(
        "Bu sertifika elektronik olarak üretilmiştir ve elle imzalanmak üzere düzenlenmiştir.\n"
        f"Oluşturulma Tarihi: {datetime.now().strftime('%d.%m.%Y %H:%M')}"
    )
    run_footer.font.size = Pt(8)
    run_footer.font.color.rgb = RGBColor(128, 128, 128)
    run_footer.font.italic = True
    
    # ================== DOSYAYI KAYDET ==================
    output_dir.mkdir(parents=True, exist_ok=True)
    filename = f"KUMPAS_Sertifika_{sertifika_no}.docx"
    file_path = output_dir / filename
    
    doc.save(str(file_path))
    
    print(f"✅ Word sertifikası oluşturuldu: {file_path}")
    return str(file_path)


if __name__ == "__main__":
    # Test için örnek veri
    test_kalibrasyon = {
        "talep_no": "TEST-001",
        "sicaklik": 20.5,
        "nem": 45,
        "uygunluk": True,
        "olcum_verileri": {
            "dis_olcum": [
                {"referans_deger": 0, "okunan_deger": 0.01, "sapma": 0.01, "belirsizlik": 0.03},
                {"referans_deger": 25, "okunan_deger": 25.02, "sapma": 0.02, "belirsizlik": 0.03},
                {"referans_deger": 50, "okunan_deger": 49.98, "sapma": -0.02, "belirsizlik": 0.03},
            ],
            "fonksiyonellik": {
                "olcme_ceneleri": "Uygun",
                "tespit_vidasi": "Uygun",
                "gosterge": "Uygun",
                "tambur_vernier": "Uygun"
            }
        }
    }
    
    test_cihaz = {
        "marka": "Mitutoyo",
        "model": "CD-15DCX",
        "seri_no": "12345678",
        "olcme_araligi": "0-150 mm"
    }
    
    test_organizasyon = {
        "musteri_adi": "Test Şirketi A.Ş.",
        "adres": "Test Mahallesi, Test Sokak No:1, İstanbul"
    }
    
    output = Path("./test_output")
    file_path = create_kumpas_certificate(
        test_kalibrasyon,
        test_cihaz,
        test_organizasyon,
        output
    )
    print(f"Test dosyası oluşturuldu: {file_path}")
